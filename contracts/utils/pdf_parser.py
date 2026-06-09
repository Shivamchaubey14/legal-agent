import os
import re
import fitz          # PyMuPDF
import pdfplumber
import pytesseract
from PIL import Image, ImageFilter, ImageEnhance
from docx import Document
import io
import logging

logger = logging.getLogger(__name__)

# ── Tesseract path (Windows) ─────────────────────────────────
pytesseract.pytesseract.tesseract_cmd = (
    r'C:\Users\Shwetdhara\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
)

# ── Contract start markers (to strip stamp paper garbage) ────
CONTRACT_START_MARKERS = [
    'TUS LEASE AGREEMENT',      # your OCR variant — most specific, check first
    'THIS LEASE AGRERMENT',
    'THIS LEASE AGREEMENT',
    'LEASE AGREEMENT',
    'THIS Contract agreement',
    'This Contract agreement',
    'THIS AGREEMENT',
    'This Agreement is made',
    'THIS DEED OF AGREEMENT',
    'This Deed of Agreement',
    'THIS DEED WITNESSETH',
    'THIS SERVICE AGREEMENT',
    'This Service Agreement',
    'EMPLOYMENT AGREEMENT',
    'Employment Agreement',
    'CONSULTANCY AGREEMENT',
    'Consultancy Agreement',
    'NOW THEREFORE',
    'WITNESSETH',
    'A. WHEREAS',               # fallback — at least gets recitals
    'WHEREAS the Lessor',
    'WHEREAS the Parties',
]


def clean_ocr_text(text: str) -> str:
    if not text:
        return text

    lines         = text.split('\n')
    cleaned_lines = []

    for line in lines:
        stripped = line.strip()

        # Skip empty lines (preserve up to 2 consecutive blanks later)
        if not stripped:
            cleaned_lines.append('')
            continue

        # ── Hard filters — always drop these ────────────────

        # Too short to be real content
        if len(stripped) < 4:
            continue

        # Lines that are mostly digits/symbols — certificate numbers,
        # barcode lines, stamp paper form fields
        digit_sym_count = sum(1 for c in stripped if c.isdigit() or c in '|/-_=+*#@[]{}\\<>')
        if digit_sym_count / len(stripped) > 0.5:
            continue

        # Lines with no letters at all
        if not any(c.isalpha() for c in stripped):
            continue

        # Certificate/reference number patterns
        import re
        if re.match(r'^[A-Z]{2}-[A-Z0-9]{10,}', stripped):
            continue
        if re.match(r'^SUBIN-', stripped):
            continue

        # Lines that look like form fields (e.g. "Certificate No. : IN-UP...")
        if re.match(r'^(Certificate|Account|Unique|Purchased|Description|Consideration|'
                    r'Stamp Duty|First Party|Second Party|Signature|Mobile No)', stripped):
            continue

        # ── Soft filter — require 50% readable chars ─────────
        alpha_space = sum(1 for c in stripped if c.isalnum() or c in ' .,;:()\'-"/')
        ratio = alpha_space / len(stripped)
        if ratio < 0.50:
            continue

        cleaned_lines.append(line)

    text = '\n'.join(cleaned_lines)

    # Collapse 3+ blank lines to max 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Collapse multiple spaces
    text = re.sub(r'[ \t]{2,}', ' ', text)
    # Remove lines that are just punctuation/dashes
    text = re.sub(r'\n[^a-zA-Z\n]{1,5}\n', '\n', text)

    return text.strip()


def strip_stamp_paper_header(text: str) -> str:
    if not text:
        return text

    for marker in CONTRACT_START_MARKERS:
        idx = text.find(marker)
        if idx > 100:          # ← lowered from 300 to 100
            logger.info(f'Stripped {idx} chars at marker: "{marker}"')
            return text[idx:]

    return text


class ContractParser:
    """
    Parses PDF and DOCX contracts into clean text.
    Strategy:
      1. Try PyMuPDF (fast, text-based PDFs)
      2. Fall back to pdfplumber
      3. Fall back to Tesseract OCR with image preprocessing (scanned PDFs)
    After parsing, strips stamp paper garbage header and cleans OCR artifacts.
    """

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.ext       = os.path.splitext(file_path)[1].lower()

    # Max pages we'll attempt OCR on — avoids timeout on giant scans
    MAX_OCR_PAGES = 60

    def parse(self) -> dict:
        """
        Returns:
            {
                'text':       str,
                'page_count': int,
                'method':     str,
                'error':      str | None
            }
        """
        # ── Pre-flight checks ────────────────────────────────
        if not os.path.exists(self.file_path):
            return self._error('File not found on disk.')

        file_size = os.path.getsize(self.file_path)
        if file_size == 0:
            return self._error('File is empty (0 bytes).')

        if file_size > 25 * 1024 * 1024:   # 25 MB hard cap
            return self._error('File exceeds 25 MB limit.')

        if self.ext not in ('.pdf', '.docx'):
            return self._error(f'Unsupported file type: {self.ext}. Only PDF and DOCX are allowed.')

        try:
            if self.ext == '.docx':
                return self._parse_docx()
            elif self.ext == '.pdf':
                return self._parse_pdf()
        except MemoryError:
            logger.error(f'MemoryError parsing {self.file_path}')
            return self._error('File is too large to process — ran out of memory.')
        except Exception as e:
            logger.error(f'ContractParser unexpected error for {self.file_path}: {e}')
            return self._error(f'Unexpected error during parsing: {e}')

    # ── PDF ──────────────────────────────────────────────────
    def _parse_pdf(self) -> dict:
        # Step 1: PyMuPDF
        result = self._try_pymupdf()
        if result and self._is_usable(result['text']):
            result['text'] = self._post_process(result['text'])
            return result

        # Step 2: pdfplumber
        result = self._try_pdfplumber()
        if result and self._is_usable(result['text']):
            result['text'] = self._post_process(result['text'])
            return result

        # Step 3: OCR fallback with preprocessing
        result = self._try_ocr()
        if result and result['text']:
            result['text'] = self._post_process(result['text'])
        return result

    def _try_pymupdf(self) -> dict | None:
        try:
            doc = fitz.open(self.file_path)
        except fitz.FileDataError as e:
            logger.warning(f'PyMuPDF: corrupt or unreadable PDF — {e}')
            return None
        except Exception as e:
            logger.warning(f'PyMuPDF failed to open: {e}')
            return None

        try:
            pages      = []
            page_count = len(doc)

            if page_count == 0:
                doc.close()
                logger.warning('PyMuPDF: PDF has 0 pages')
                return None

            for page in doc:
                try:
                    text = page.get_text('text')
                    pages.append(text)
                except Exception as page_err:
                    logger.warning(f'PyMuPDF: error on page {page.number}: {page_err}')
                    pages.append('')   # keep page slot, skip broken page

            doc.close()
            full_text = '\n\n'.join(pages).strip()

            return {
                'text':       full_text,
                'page_count': page_count,
                'method':     'pymupdf',
                'error':      None,
            }
        except Exception as e:
            logger.warning(f'PyMuPDF failed during extraction: {e}')
            try: doc.close()
            except: pass
            return None

    def _try_pdfplumber(self) -> dict | None:
        try:
            pages = []
            with pdfplumber.open(self.file_path) as pdf:
                page_count = len(pdf.pages)

                if page_count == 0:
                    logger.warning('pdfplumber: PDF has 0 pages')
                    return None

                for i, page in enumerate(pdf.pages):
                    try:
                        text = page.extract_text() or ''
                        pages.append(text)
                    except Exception as page_err:
                        logger.warning(f'pdfplumber: error on page {i}: {page_err}')
                        pages.append('')

            full_text = '\n\n'.join(pages).strip()
            return {
                'text':       full_text,
                'page_count': page_count,
                'method':     'pdfplumber',
                'error':      None,
            }
        except pdfplumber.utils.exceptions.PDFSyntaxError as e:
            logger.warning(f'pdfplumber: PDF syntax error (corrupt file) — {e}')
            return None
        except Exception as e:
            logger.warning(f'pdfplumber failed: {e}')
            return None

    def _try_ocr(self) -> dict:
        try:
            doc        = fitz.open(self.file_path)
            page_count = len(doc)
            pages      = []

            if len(doc) > self.MAX_OCR_PAGES:
                logger.warning(f'OCR: truncating to {self.MAX_OCR_PAGES} pages (file has {len(doc)})')

            for page_num, page in enumerate(doc):
                if page_num >= self.MAX_OCR_PAGES:
                    break
                # Render at 400 DPI instead of 300 — huge quality boost
                mat = fitz.Matrix(400 / 72, 400 / 72)
                pix = page.get_pixmap(matrix=mat)
                img = Image.open(io.BytesIO(pix.tobytes('png')))

                # ── Aggressive preprocessing ─────────────────────
                img = img.convert('L')                        # grayscale

                # Resize up 1.5x before binarizing — helps Tesseract a lot
                w, h = img.size
                img = img.resize((int(w * 1.5), int(h * 1.5)), Image.LANCZOS)

                # Strong contrast
                enhancer = ImageEnhance.Contrast(img)
                img      = enhancer.enhance(2.5)

                # Sharpen twice
                img = img.filter(ImageFilter.SHARPEN)
                img = img.filter(ImageFilter.SHARPEN)

                # Denoise
                img = img.filter(ImageFilter.MedianFilter(size=3))

                # Binarize with a higher threshold — stamp paper bg is dark
                img = img.point(lambda x: 0 if x < 160 else 255, '1')

                # OCR — psm 6 = assume uniform block of text
                custom_config = r'--oem 3 --psm 6 -l eng+hin'
                text = pytesseract.image_to_string(img, config=custom_config)
                pages.append(text)
                logger.info(f'OCR page {page_num + 1}/{page_count}: {len(text)} chars')

            doc.close()
            full_text = '\n\n'.join(pages).strip()
            return {'text': full_text, 'page_count': page_count, 'method': 'ocr', 'error': None} 

        except MemoryError:
            logger.error('OCR failed: out of memory — file may be too large')
            return self._error('OCR ran out of memory. Try a smaller file.')
        except Exception as e:
            logger.error(f'OCR failed: {e}')
            return self._error(f'OCR failed: {e}')

    # ── DOCX ─────────────────────────────────────────────────
    def _parse_docx(self) -> dict:
        try:
            doc        = Document(self.file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())

            full_text  = '\n\n'.join(paragraphs).strip()
            page_count = max(1, len(full_text) // 3000)  # Estimate

            return {
                'text':       full_text,
                'page_count': page_count,
                'method':     'docx',
                'error':      None,
            }
        except Exception as e:
            logger.error(f'DOCX parsing failed: {e}')
            # Common cause: password-protected or corrupt DOCX
            if 'encrypted' in str(e).lower() or 'password' in str(e).lower():
                return self._error('This DOCX file is password-protected. Please remove the password and re-upload.')
            if 'BadZipFile' in type(e).__name__ or 'zipfile' in str(e).lower():
                return self._error('This DOCX file appears to be corrupt or is not a valid Word document.')
            return self._error(f'Could not read DOCX file: {e}')

    # ── Post processing ──────────────────────────────────────
    def _post_process(self, text: str) -> str:
        """
        Run all cleaning steps after extraction:
        1. Clean OCR garbage lines
        2. Strip stamp paper header
        """
        text = clean_ocr_text(text)
        text = strip_stamp_paper_header(text)
        return text

    # ── Helpers ──────────────────────────────────────────────
    @staticmethod
    def _is_usable(text: str, min_chars: int = 100) -> bool:
        """Check if extracted text is actually useful."""
        if not text:
            return False
        clean = text.replace('\n', '').replace(' ', '')
        return len(clean) >= min_chars

    @staticmethod
    def _error(msg: str) -> dict:
        return {
            'text':       '',
            'page_count': 0,
            'method':     'failed',
            'error':      msg,
        }


def parse_contract_file(file_path: str) -> dict:
    """Convenience function — call this from views/tasks."""
    parser = ContractParser(file_path)
    return parser.parse()